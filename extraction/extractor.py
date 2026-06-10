"""
extraction/extractor.py

Clean extraction layer that converts PDF / DOCX / TXT files into the
unified Document model.

Strategy
--------
* **Primary path** — Docling's `DocumentConverter` for PDF and DOCX.
  Docling preserves structural semantics (headings, tables, lists) far
  better than raw text extraction.
* **TXT path** — heuristic line-based parser (no Docling needed).
* **Fallback** — if Docling is unavailable or fails, we fall back to
  python-docx (DOCX) or pdfplumber (PDF) for plain-text extraction.

The result is always a `Document` model, so downstream stages are
completely decoupled from the file format.
"""

from __future__ import annotations

import io
import logging
import re
import tempfile
from pathlib import Path
from typing import BinaryIO

from models import (
    Document,
    DocumentMetadata,
    DocumentSection,
    Table,
)
from config import CONFIG

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _word_count(text: str) -> int:
    return len(text.split())


def _build_metadata(filename: str, file_type: str, doc: Document) -> DocumentMetadata:
    full = doc.full_text()
    return DocumentMetadata(
        filename=filename,
        file_type=file_type,
        word_count=_word_count(full),
        char_count=len(full),
    )


# ---------------------------------------------------------------------------
# TXT extractor (heuristic)
# ---------------------------------------------------------------------------

def _extract_txt(content: bytes, filename: str) -> Document:
    text = content.decode(CONFIG.extraction.txt_encoding, errors="replace")
    lines = text.splitlines()

    sections: list[DocumentSection] = []
    current = DocumentSection()
    buffer: list[str] = []

    HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)|^([A-Z][A-Z0-9 ]{3,}):?\s*$")

    def flush_buffer():
        nonlocal buffer
        para = " ".join(buffer).strip()
        if para:
            current.paragraphs.append(para)
        buffer = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush_buffer()
            continue

        m = HEADING_RE.match(stripped)
        if m:
            flush_buffer()
            # Save previous section
            if current.paragraphs or current.heading:
                sections.append(current)
            current = DocumentSection()
            if m.group(1):
                current.heading = m.group(2).strip()
                current.heading_level = len(m.group(1))
            else:
                current.heading = m.group(3).strip()
                current.heading_level = 1
        elif stripped.startswith(("- ", "* ", "• ")):
            flush_buffer()
            item = stripped[2:].strip()
            if current.lists:
                current.lists[-1].append(item)
            else:
                current.lists.append([item])
        else:
            buffer.append(stripped)

    flush_buffer()
    if current.paragraphs or current.heading or current.lists:
        sections.append(current)

    if not sections:
        sections.append(DocumentSection(paragraphs=[text.strip()]))

    doc = Document(sections=sections)
    doc.metadata = _build_metadata(filename, "txt", doc)
    return doc


# ---------------------------------------------------------------------------
# Docling-based extractor (PDF & DOCX)
# ---------------------------------------------------------------------------

def _docling_to_document(docling_doc, filename: str, file_type: str) -> Document:
    """Convert a Docling DoclingDocument into our Document model."""
    sections: list[DocumentSection] = []
    current = DocumentSection()

    try:
        # Docling's export_to_markdown gives us a structured string;
        # we use the element iterator for fine-grained structure.
        from docling.datamodel.base_models import InputFormat  # noqa
        from docling_core.types.doc import DocItemLabel

        for item, _level in docling_doc.iterate_items():
            label = getattr(item, "label", None)
            text = ""
            if hasattr(item, "text"):
                text = item.text or ""
            elif hasattr(item, "export_to_markdown"):
                text = item.export_to_markdown()

            text = text.strip()
            if not text:
                continue

            if label in (DocItemLabel.SECTION_HEADER, DocItemLabel.TITLE):
                if current.paragraphs or current.heading or current.tables:
                    sections.append(current)
                current = DocumentSection(heading=text, heading_level=1)

            elif label == DocItemLabel.LIST_ITEM:
                if current.lists:
                    current.lists[-1].append(text)
                else:
                    current.lists.append([text])

            elif label == DocItemLabel.TABLE:
                # Try to extract table structure
                tbl = Table()
                if hasattr(item, "data") and item.data:
                    grid = item.data.grid
                    if grid:
                        first_row = [c.text for c in grid[0]]
                        tbl.headers = first_row
                        for row in grid[1:]:
                            tbl.rows.append([c.text for c in row])
                else:
                    tbl = Table(caption=text)
                current.tables.append(tbl)

            else:
                current.paragraphs.append(text)

    except Exception:
        # Fallback: export whole doc as markdown, parse as TXT
        md = docling_doc.export_to_markdown()
        return _extract_txt(md.encode(), filename)

    if current.paragraphs or current.heading or current.tables or current.lists:
        sections.append(current)

    if not sections:
        md = docling_doc.export_to_markdown()
        return _extract_txt(md.encode(), filename)

    # Missing pages detection
    all_pages = []
    if hasattr(docling_doc, "pages") and isinstance(docling_doc.pages, dict):
        all_pages = sorted(list(docling_doc.pages.keys()))
    
    missing_pages = []
    if all_pages:
        min_p, max_p = min(all_pages), max(all_pages)
        missing_pages = [p for p in range(min_p, max_p + 1) if p not in all_pages]

    doc = Document(sections=sections)
    doc.metadata = _build_metadata(filename, file_type, doc)
    doc.metadata.missing_pages = missing_pages
    return doc


def _extract_with_docling(path: Path, filename: str, file_type: str) -> Document:
    from docling.document_converter import DocumentConverter
    converter = DocumentConverter()
    result = converter.convert(str(path))
    return _docling_to_document(result.document, filename, file_type)


# ---------------------------------------------------------------------------
# Fallback extractors
# ---------------------------------------------------------------------------

def _extract_pdf_fallback(content: bytes, filename: str) -> Document:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        text = "\n\n".join(pages)
        doc = _extract_txt(text.encode(), filename)
        doc.metadata.file_type = "pdf"
        doc.metadata.page_count = len(pages)
        return doc
    except ImportError:
        # Last resort: pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n\n".join(pages)
            doc = _extract_txt(text.encode(), filename)
            doc.metadata.file_type = "pdf"
            doc.metadata.page_count = len(pages)
            return doc
        except Exception as exc:
            raise RuntimeError(f"PDF extraction failed: {exc}") from exc


def _extract_docx_fallback(content: bytes, filename: str) -> Document:
    import docx  # python-docx
    document = docx.Document(io.BytesIO(content))

    sections: list[DocumentSection] = []
    current = DocumentSection()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style:
            if current.paragraphs or current.heading:
                sections.append(current)
            level = int(style[-1]) if style[-1].isdigit() else 1
            current = DocumentSection(heading=text, heading_level=level)
        else:
            current.paragraphs.append(text)

    if current.paragraphs or current.heading:
        sections.append(current)

    for table in document.tables:
        tbl = Table()
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if i == 0:
                tbl.headers = cells
            else:
                tbl.rows.append(cells)
        if sections:
            sections[-1].tables.append(tbl)
        else:
            s = DocumentSection()
            s.tables.append(tbl)
            sections.append(s)

    if not sections:
        sections.append(DocumentSection(paragraphs=[document.core_properties.description or ""]))

    doc = Document(sections=sections)
    doc.metadata = _build_metadata(filename, "docx", doc)
    return doc


# ---------------------------------------------------------------------------
# Public extractor class
# ---------------------------------------------------------------------------

class DocumentExtractor:
    """
    Entry point for document extraction.

    Usage
    -----
    extractor = DocumentExtractor()
    doc = extractor.extract(file_bytes, "report.pdf")
    """

    def extract(self, content: bytes, filename: str) -> Document:
        """
        Extract a Document from raw bytes.

        Parameters
        ----------
        content : bytes
            Raw file bytes (PDF, DOCX, or TXT).
        filename : str
            Original filename — used to infer the file type.

        Returns
        -------
        Document
        """
        suffix = Path(filename).suffix.lower()

        if suffix == ".txt":
            return _extract_txt(content, filename)

        # For PDF / DOCX: prefer Docling, fall back gracefully
        try:
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(content)
                tmp_path = Path(tmp.name)
            return _extract_with_docling(tmp_path, filename, suffix.lstrip("."))
        except Exception as exc:
            logger.warning("Docling extraction failed (%s), using fallback: %s", filename, exc)
            if suffix == ".pdf":
                return _extract_pdf_fallback(content, filename)
            elif suffix in (".docx", ".doc"):
                return _extract_docx_fallback(content, filename)
            else:
                return _extract_txt(content, filename)
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass

    def extract_from_uploaded(self, uploaded_file) -> Document:
        """
        Convenience wrapper for Streamlit UploadedFile objects.
        """
        return self.extract(uploaded_file.read(), uploaded_file.name)
